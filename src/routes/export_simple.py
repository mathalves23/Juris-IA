from flask import Blueprint, request, jsonify, send_file, make_response
from flask_jwt_extended import jwt_required, get_jwt_identity
import os
import tempfile
from datetime import datetime
from io import BytesIO
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.extensions import db
from src.models.document import Document
from src.models.template import Template

export_bp = Blueprint('export', __name__)

def replace_variables(content, variables):
    """Substitui variáveis no conteúdo"""
    if not variables:
        return content
    
    for key, value in variables.items():
        placeholder = f"{{{key}}}"
        content = content.replace(placeholder, str(value))
    
    return content

@export_bp.route('/document/<int:document_id>/pdf', methods=['GET'])
@jwt_required()
def export_document_pdf(document_id):
    """Exporta documento como PDF (versão simplificada)"""
    try:
        user_id = get_jwt_identity()
        document = Document.query.filter_by(id=document_id, user_id=user_id).first()
        
        if not document:
            return jsonify({"error": "Documento não encontrado"}), 404
        
        # Por enquanto, retorna uma mensagem informando sobre a limitação
        return jsonify({
            "message": "Exportação PDF temporariamente indisponível. Use a exportação DOCX.",
            "alternative": f"/api/export/document/{document_id}/docx"
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Erro ao exportar PDF: {str(e)}"}), 500

@export_bp.route('/document/<int:document_id>/docx', methods=['GET'])
@jwt_required()
def export_document_docx(document_id):
    """Exporta documento como DOCX"""
    try:
        user_id = get_jwt_identity()
        document = Document.query.filter_by(id=document_id, user_id=user_id).first()
        
        if not document:
            return jsonify({"error": "Documento não encontrado"}), 404
        
        # Criar documento Word
        doc = docx.Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Adicionar título
        title = doc.add_heading(document.titulo, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar conteúdo
        # Remove tags HTML básicas
        content = document.conteudo
        content = content.replace('<p>', '\n').replace('</p>', '')
        content = content.replace('<br>', '\n').replace('<br/>', '\n')
        content = content.replace('<strong>', '').replace('</strong>', '')
        content = content.replace('<em>', '').replace('</em>', '')
        
        # Dividir em parágrafos
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Adicionar rodapé com data
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salvar em memória
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Preparar resposta
        filename = f"{document.titulo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        response = make_response(file_stream.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        return jsonify({"error": f"Erro ao exportar DOCX: {str(e)}"}), 500

@export_bp.route('/template/<int:template_id>/pdf', methods=['POST'])
@jwt_required()
def export_template_pdf(template_id):
    """Exporta template preenchido como PDF (versão simplificada)"""
    try:
        user_id = get_jwt_identity()
        template = Template.query.filter_by(id=template_id, user_id=user_id).first()
        
        if not template:
            return jsonify({"error": "Template não encontrado"}), 404
        
        # Por enquanto, retorna uma mensagem informando sobre a limitação
        return jsonify({
            "message": "Exportação PDF temporariamente indisponível. Use a exportação DOCX.",
            "alternative": f"/api/export/template/{template_id}/docx"
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Erro ao exportar PDF: {str(e)}"}), 500

@export_bp.route('/template/<int:template_id>/docx', methods=['POST'])
@jwt_required()
def export_template_docx(template_id):
    """Exporta template preenchido como DOCX"""
    try:
        user_id = get_jwt_identity()
        template = Template.query.filter_by(id=template_id, user_id=user_id).first()
        
        if not template:
            return jsonify({"error": "Template não encontrado"}), 404
        
        data = request.get_json()
        variables = data.get('variables', {})
        
        # Substituir variáveis no conteúdo
        content = replace_variables(template.conteudo, variables)
        
        # Criar documento Word
        doc = docx.Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Adicionar título
        title = doc.add_heading(template.titulo, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar conteúdo
        # Remove tags HTML básicas
        content = content.replace('<p>', '\n').replace('</p>', '')
        content = content.replace('<br>', '\n').replace('<br/>', '\n')
        content = content.replace('<strong>', '').replace('</strong>', '')
        content = content.replace('<em>', '').replace('</em>', '')
        
        # Dividir em parágrafos
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Adicionar rodapé com data
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salvar em memória
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Preparar resposta
        filename = f"{template.titulo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        response = make_response(file_stream.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        return jsonify({"error": f"Erro ao exportar DOCX: {str(e)}"}), 500 