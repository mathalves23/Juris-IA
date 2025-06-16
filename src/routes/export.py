from flask import Blueprint, request, jsonify, send_file, make_response
from flask_jwt_extended import jwt_required, get_jwt_identity
import os
import tempfile
from datetime import datetime
from io import BytesIO
import docx
from weasyprint import HTML, CSS
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.extensions import db
from src.models.document import Document

export_bp = Blueprint('export', __name__)

def html_to_clean_text(html_content):
    """Converte HTML para texto limpo removendo tags"""
    import re
    # Remove tags HTML
    clean = re.compile('<.*?>')
    text = re.sub(clean, '', html_content)
    # Decodifica entidades HTML
    import html
    text = html.unescape(text)
    return text

def create_pdf_from_html(content, title="Documento"):
    """Cria PDF a partir de conteúdo HTML"""
    html_template = f"""
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title}</title>
        <style>
            @page {{
                margin: 2cm;
                size: A4;
                @bottom-center {{
                    content: "Página " counter(page) " de " counter(pages);
                    font-size: 10px;
                    color: #666;
                }}
            }}
            body {{
                font-family: 'Times New Roman', serif;
                font-size: 12px;
                line-height: 1.6;
                color: #333;
                margin: 0;
                padding: 0;
            }}
            h1, h2, h3 {{
                color: #000;
                margin-top: 1.5em;
                margin-bottom: 0.5em;
            }}
            h1 {{
                font-size: 18px;
                text-align: center;
                font-weight: bold;
            }}
            h2 {{
                font-size: 16px;
                font-weight: bold;
            }}
            h3 {{
                font-size: 14px;
                font-weight: bold;
            }}
            p {{
                margin-bottom: 1em;
                text-align: justify;
            }}
            .header {{
                text-align: center;
                margin-bottom: 2em;
                border-bottom: 1px solid #ccc;
                padding-bottom: 1em;
            }}
            .content {{
                margin-top: 1em;
            }}
            strong {{
                font-weight: bold;
            }}
            em {{
                font-style: italic;
            }}
            ul, ol {{
                margin-left: 2em;
                margin-bottom: 1em;
            }}
            li {{
                margin-bottom: 0.5em;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{title}</h1>
            <p style="font-size: 10px; color: #666;">
                Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}
            </p>
        </div>
        <div class="content">
            {content}
        </div>
    </body>
    </html>
    """
    
    # Gerar PDF
    pdf_file = BytesIO()
    HTML(string=html_template).write_pdf(pdf_file)
    pdf_file.seek(0)
    return pdf_file

def create_docx_from_text(content, title="Documento"):
    """Cria DOCX a partir de conteúdo texto"""
    doc = docx.Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Adicionar título
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar data de geração
    date_para = doc.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar quebra de linha
    doc.add_paragraph()
    
    # Converter HTML para texto limpo
    clean_text = html_to_clean_text(content)
    
    # Dividir em parágrafos e adicionar ao documento
    paragraphs = clean_text.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Salvar em BytesIO
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

@export_bp.route('/document/<int:document_id>/pdf', methods=['GET'])
@jwt_required()
def export_document_pdf(document_id):
    """Exporta documento como PDF"""
    usuario_id = get_jwt_identity()
    
    # Buscar documento
    document = Document.query.get(document_id)
    if not document:
        return jsonify({"error": "Documento não encontrado"}), 404
    
    # Verificar se o usuário é o proprietário
    if document.usuario_id != usuario_id:
        return jsonify({"error": "Acesso negado"}), 403
    
    try:
        # Gerar PDF
        pdf_file = create_pdf_from_html(document.conteudo, document.titulo)
        
        # Criar nome do arquivo
        filename = f"{document.titulo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Retornar arquivo
        response = make_response(pdf_file.read())
        response.headers.set('Content-Type', 'application/pdf')
        response.headers.set('Content-Disposition', f'attachment; filename="{filename}"')
        return response
        
    except Exception as e:
        return jsonify({"error": f"Erro ao gerar PDF: {str(e)}"}), 500

@export_bp.route('/document/<int:document_id>/docx', methods=['GET'])
@jwt_required()
def export_document_docx(document_id):
    """Exporta documento como DOCX"""
    usuario_id = get_jwt_identity()
    
    # Buscar documento
    document = Document.query.get(document_id)
    if not document:
        return jsonify({"error": "Documento não encontrado"}), 404
    
    # Verificar se o usuário é o proprietário
    if document.usuario_id != usuario_id:
        return jsonify({"error": "Acesso negado"}), 403
    
    try:
        # Gerar DOCX
        docx_file = create_docx_from_text(document.conteudo, document.titulo)
        
        # Criar nome do arquivo
        filename = f"{document.titulo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Retornar arquivo
        response = make_response(docx_file.read())
        response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response.headers.set('Content-Disposition', f'attachment; filename="{filename}"')
        return response
        
    except Exception as e:
        return jsonify({"error": f"Erro ao gerar DOCX: {str(e)}"}), 500

@export_bp.route('/template/<int:template_id>/pdf', methods=['POST'])
@jwt_required()
def export_template_pdf(template_id):
    """Exporta template preenchido como PDF"""
    usuario_id = get_jwt_identity()
    
    from src.models.template import Template
    
    # Buscar template
    template = Template.query.get(template_id)
    if not template:
        return jsonify({"error": "Template não encontrado"}), 404
    
    # Verificar se o usuário tem acesso
    if template.usuario_id != usuario_id and not template.publico:
        return jsonify({"error": "Acesso negado"}), 403
    
    try:
        data = request.get_json()
        variables = data.get('variables', {})
        
        # Aplicar variáveis ao conteúdo
        content = template.conteudo
        for var_name, var_value in variables.items():
            content = content.replace(f"{{{var_name}}}", var_value)
        
        # Gerar PDF
        pdf_file = create_pdf_from_html(content, template.titulo)
        
        # Criar nome do arquivo
        filename = f"{template.titulo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Retornar arquivo
        response = make_response(pdf_file.read())
        response.headers.set('Content-Type', 'application/pdf')
        response.headers.set('Content-Disposition', f'attachment; filename="{filename}"')
        return response
        
    except Exception as e:
        return jsonify({"error": f"Erro ao gerar PDF: {str(e)}"}), 500

@export_bp.route('/template/<int:template_id>/docx', methods=['POST'])
@jwt_required()
def export_template_docx(template_id):
    """Exporta template preenchido como DOCX"""
    usuario_id = get_jwt_identity()
    
    from src.models.template import Template
    
    # Buscar template
    template = Template.query.get(template_id)
    if not template:
        return jsonify({"error": "Template não encontrado"}), 404
    
    # Verificar se o usuário tem acesso
    if template.usuario_id != usuario_id and not template.publico:
        return jsonify({"error": "Acesso negado"}), 403
    
    try:
        data = request.get_json()
        variables = data.get('variables', {})
        
        # Aplicar variáveis ao conteúdo
        content = template.conteudo
        for var_name, var_value in variables.items():
            content = content.replace(f"{{{var_name}}}", var_value)
        
        # Gerar DOCX
        docx_file = create_docx_from_text(content, template.titulo)
        
        # Criar nome do arquivo
        filename = f"{template.titulo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Retornar arquivo
        response = make_response(docx_file.read())
        response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response.headers.set('Content-Disposition', f'attachment; filename="{filename}"')
        return response
        
    except Exception as e:
        return jsonify({"error": f"Erro ao gerar DOCX: {str(e)}"}), 500 